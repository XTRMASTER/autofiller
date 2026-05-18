import docx
import os

class WordHandler:
    @staticmethod
    def scan_word_document(filepath: str):
        """
        Scans a Word document and returns its paragraphs and text content.
        This is a basic implementation that reads text from paragraphs and tables.
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        doc = docx.Document(filepath)
        content = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content.append({'type': 'paragraph', 'index': i, 'text': para.text})

        for i, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        content.append({
                            'type': 'table_cell',
                            'table_idx': i,
                            'row_idx': row_idx,
                            'col_idx': col_idx,
                            'text': cell.text
                        })
        return content

    @staticmethod
    def apply_word_variables(filepath: str, variables_dict: dict, output_path: str):
        """
        Replaces variables in a Word document and saves to output_path.
        variables_dict maps text to replace to the new value.
        Preserves simple formatting by replacing text within runs.
        """
        doc = docx.Document(filepath)

        # Replace in paragraphs
        for para in doc.paragraphs:
            WordHandler._replace_in_paragraph(para, variables_dict)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        WordHandler._replace_in_paragraph(para, variables_dict)

        doc.save(output_path)

    @staticmethod
    def _replace_in_paragraph(para, variables_dict):
        for old_val, new_val in variables_dict.items():
            start_search = 0
            while True:
                idx = WordHandler._replace_text_in_paragraph(para, old_val, new_val, start_search)
                if idx == -1:
                    break
                start_search = idx + len(new_val)

    @staticmethod
    def _replace_text_in_paragraph(paragraph, old_text, new_text, start_search=0):
        # We need the full text to find the match correctly considering start_search
        text = ""
        for run in paragraph.runs:
            text += run.text

        start_idx = text.find(old_text, start_search)
        if start_idx == -1:
            return -1

        end_idx = start_idx + len(old_text)

        # Check if it fits entirely within one run
        curr_len = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            if curr_len <= start_idx and curr_len + run_len >= end_idx:
                # The match is entirely within this single run
                run_start = start_idx - curr_len
                run_end = end_idx - curr_len
                run.text = run.text[:run_start] + new_text + run.text[run_end:]
                return start_idx
            curr_len += run_len

        # It spans multiple runs
        # Find start run
        curr_len = 0
        start_run_idx = -1
        start_run_offset = -1
        for i, run in enumerate(paragraph.runs):
            if curr_len + len(run.text) > start_idx:
                start_run_idx = i
                start_run_offset = start_idx - curr_len
                break
            curr_len += len(run.text)

        # Find end run
        curr_len = 0
        end_run_idx = -1
        end_run_offset = -1
        for i, run in enumerate(paragraph.runs):
            if curr_len + len(run.text) >= end_idx:
                end_run_idx = i
                end_run_offset = end_idx - curr_len
                break
            curr_len += len(run.text)

        # Now modify runs
        if start_run_idx != -1 and end_run_idx != -1:
            start_run_text = paragraph.runs[start_run_idx].text
            end_run_text = paragraph.runs[end_run_idx].text

            if start_run_idx == end_run_idx:
                paragraph.runs[start_run_idx].text = start_run_text[:start_run_offset] + new_text + start_run_text[end_run_offset:]
            else:
                # Put new text in start run
                paragraph.runs[start_run_idx].text = start_run_text[:start_run_offset] + new_text

                # Clear intermediate runs
                for i in range(start_run_idx + 1, end_run_idx):
                    paragraph.runs[i].text = ""

                # Adjust end run
                paragraph.runs[end_run_idx].text = end_run_text[end_run_offset:]

        return start_idx
