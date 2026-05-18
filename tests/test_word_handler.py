import os
import unittest
import docx
from app.core.word_handler import WordHandler

class TestWordHandler(unittest.TestCase):
    def setUp(self):
        self.test_dir = os.path.dirname(os.path.abspath(__file__))
        self.test_file = os.path.join(self.test_dir, 'test_preserve.docx')
        self.output_file = os.path.join(self.test_dir, 'test_preserve_out.docx')

        # Create a test document with complex formatting
        doc = docx.Document()
        p = doc.add_paragraph()
        r1 = p.add_run('This is a t')
        r2 = p.add_run('es')
        r2.bold = True
        r3 = p.add_run('t of ')
        r4 = p.add_run('{{v')
        r4.italic = True
        r5 = p.add_run('ar}} ')
        r5.underline = True
        r6 = p.add_run('with formatting.')
        doc.save(self.test_file)

    def tearDown(self):
        if os.path.exists(self.test_file):
            os.remove(self.test_file)
        if os.path.exists(self.output_file):
            os.remove(self.output_file)

    def test_preserve_formatting(self):
        variables = {'test': 'REPLACEMENT', '{{var}}': 'VALUE'}
        WordHandler.apply_word_variables(self.test_file, variables, self.output_file)

        doc = docx.Document(self.output_file)
        p = doc.paragraphs[0]
        self.assertEqual(p.text, 'This is a REPLACEMENT of VALUE with formatting.')

        # Check that runs were preserved correctly
        runs = p.runs
        self.assertEqual(runs[0].text, 'This is a REPLACEMENT')
        self.assertIsNone(runs[0].bold)

        # The bold 'es' run was cleared, so it should be empty text, but still exist and be bold
        self.assertEqual(runs[1].text, '')
        self.assertTrue(runs[1].bold)

        self.assertEqual(runs[2].text, ' of ')
        self.assertIsNone(runs[2].bold)

        self.assertEqual(runs[3].text, 'VALUE')
        self.assertTrue(runs[3].italic)

        self.assertEqual(runs[4].text, ' ')
        self.assertTrue(runs[4].underline)

        self.assertEqual(runs[5].text, 'with formatting.')

if __name__ == '__main__':
    unittest.main()
