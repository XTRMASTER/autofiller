import os
import unittest
import docx
from app.core.word_handler import WordHandler

class TestWordHandler(unittest.TestCase):
    def setUp(self):
        self.test_dir = os.path.dirname(os.path.abspath(__file__))
        self.test_file = os.path.join(self.test_dir, 'test_rec.docx')
        self.output_file = os.path.join(self.test_dir, 'test_rec_out.docx')

        doc = docx.Document()
        p = doc.add_paragraph()
        r1 = p.add_run('This is a test.')
        doc.save(self.test_file)

    def tearDown(self):
        if os.path.exists(self.test_file):
            os.remove(self.test_file)
        if os.path.exists(self.output_file):
            os.remove(self.output_file)

    def test_recursion(self):
        variables = {'test': 'test case'}
        # This will trigger recursion error if the fix is not applied
        WordHandler.apply_word_variables(self.test_file, variables, self.output_file)

        doc = docx.Document(self.output_file)
        p = doc.paragraphs[0]
        self.assertEqual(p.text, 'This is a test case.')

if __name__ == '__main__':
    unittest.main()
